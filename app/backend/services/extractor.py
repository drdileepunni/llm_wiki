import hashlib
import logging
import fitz  # PyMuPDF
import httpx
from pathlib import Path
from urllib.parse import urljoin, urlparse

log = logging.getLogger("wiki.extractor")

# Safety ceiling — vision-transcribing more than this many pages per PDF is
# very slow and expensive.  Pages beyond the limit are extracted with PyMuPDF
# text fallback and a warning is added.
MAX_VISION_PAGES = 150


# ── PDF page cache ────────────────────────────────────────────────────────────

class PdfPageCache:
    """
    Disk-based cache for vision-transcribed PDF pages.

    Layout:  <CACHE_DIR>/pdf_pages/<file_hash>/page_<NNNN>.txt

    The cache key is the first 20 hex chars of the SHA-256 of the PDF bytes,
    so the same file re-uploaded under a different name still gets a cache hit.
    Cache entries are plain UTF-8 text files — trivially inspectable.
    """

    def __init__(self, cache_dir: "Path | None" = None):
        if cache_dir is None:
            from ..config import CACHE_DIR
            cache_dir = CACHE_DIR
        self._root = cache_dir / "pdf_pages"
        self._root.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def file_hash(path: Path) -> str:
        sha = hashlib.sha256()
        with open(path, "rb") as f:
            for block in iter(lambda: f.read(65536), b""):
                sha.update(block)
        return sha.hexdigest()[:20]

    def _page_path(self, file_hash: str, page_num: int) -> Path:
        doc_dir = self._root / file_hash
        doc_dir.mkdir(exist_ok=True)
        return doc_dir / f"page_{page_num:04d}.txt"

    def get(self, file_hash: str, page_num: int) -> str | None:
        p = self._page_path(file_hash, page_num)
        return p.read_text(encoding="utf-8") if p.exists() else None

    def put(self, file_hash: str, page_num: int, text: str) -> None:
        self._page_path(file_hash, page_num).write_text(text, encoding="utf-8")

    def cached_pages(self, file_hash: str) -> set[int]:
        """Return the set of already-cached page numbers for a given file hash."""
        doc_dir = self._root / file_hash
        if not doc_dir.exists():
            return set()
        return {
            int(p.stem.split("_")[1])
            for p in doc_dir.glob("page_*.txt")
        }


_pdf_cache: "PdfPageCache | None" = None  # lazy singleton for default KB


def extract_pdf_text(path: Path) -> str:
    """Fast PyMuPDF text extraction — good for born-digital PDFs."""
    doc = fitz.open(str(path))
    return "\n\n".join(page.get_text() for page in doc)


def extract_pdf_vision(path: Path, cache_dir: "Path | None" = None) -> str:
    """
    Render every page of the PDF to a 144 DPI PNG and pass each one to the
    active LLM for faithful transcription.

    Why vision over plain text extraction:
      - Handles scanned / image-based PDFs
      - Preserves table structure, sidebars, callouts
      - Understands flowcharts and algorithm diagrams
      - Correct reading order even in multi-column layouts

    Costs are logged to the DB as operation="extract_pdf_vision" so they
    appear in the dashboard alongside ingest costs.
    """
    from .llm_client import get_llm_client
    from .token_tracker import log_call
    from ..config import MODEL

    global _pdf_cache
    cache = PdfPageCache(cache_dir) if cache_dir else (_pdf_cache or PdfPageCache())
    if not cache_dir and _pdf_cache is None:
        _pdf_cache = cache

    doc         = fitz.open(str(path))
    total_pages = len(doc)

    # Compute file hash once for cache lookups
    file_hash     = cache.file_hash(path)
    already_cached = cache.cached_pages(file_hash)

    log.info(
        "PDF vision extraction: %s  pages=%d  model=%s  cache_key=%s  cached_pages=%d",
        path.name, total_pages, MODEL, file_hash, len(already_cached),
    )

    if total_pages > MAX_VISION_PAGES:
        log.warning(
            "PDF has %d pages — only the first %d will use vision; "
            "remaining pages fall back to text extraction.",
            total_pages, MAX_VISION_PAGES,
        )

    llm         = get_llm_client()
    pages_text  = []
    total_in    = 0
    total_out   = 0

    # Render at 2× scale → 144 DPI for legible fine print
    render_matrix = fitz.Matrix(2.0, 2.0)

    for page_num, page in enumerate(doc, 1):

        # ── Cache hit ────────────────────────────────────────────────────────
        cached_text = cache.get(file_hash, page_num)
        if cached_text is not None:
            log.info("  Page %d/%d [cache hit]", page_num, total_pages)
            pages_text.append(f"--- Page {page_num} ---\n{cached_text}")
            continue

        # ── Beyond vision limit: text fallback ───────────────────────────────
        if page_num > MAX_VISION_PAGES:
            fallback = page.get_text()
            pages_text.append(f"--- Page {page_num} (text fallback) ---\n{fallback}")
            cache.put(file_hash, page_num, fallback)   # cache fallback too
            continue

        # ── Vision transcription ─────────────────────────────────────────────
        log.info("  Page %d/%d [transcribing]", page_num, total_pages)
        try:
            pix         = page.get_pixmap(matrix=render_matrix)
            image_bytes = pix.tobytes("png")
            page_text, usage = llm.transcribe_page(image_bytes, page_num, total_pages)
            cache.put(file_hash, page_num, page_text)   # persist immediately
            pages_text.append(f"--- Page {page_num} ---\n{page_text}")
            total_in  += usage.input_tokens
            total_out += usage.output_tokens
            log.info(
                "  Page %d done  in=%d  out=%d  chars=%d",
                page_num, usage.input_tokens, usage.output_tokens, len(page_text),
            )
        except Exception as exc:
            log.error(
                "  Vision failed for page %d (%s) — using text fallback", page_num, exc
            )
            fallback = page.get_text()
            pages_text.append(f"--- Page {page_num} (text fallback) ---\n{fallback}")
            # Don't cache failed pages — let the next attempt retry them

    # Log total extraction cost to the DB
    log_call("extract_pdf_vision", path.name, total_in, total_out, model=MODEL)
    log.info(
        "Vision extraction complete: pages=%d  total_in=%d  total_out=%d",
        total_pages, total_in, total_out,
    )

    return "\n\n".join(pages_text)

MAX_EMBEDDED_IMAGES = 15


def _transcribe_embedded_images(
    image_blobs: list[tuple[str, bytes]],   # [(label, bytes), ...]
    source_name: str,
    operation: str,
) -> str:
    """
    Vision-transcribe a list of embedded images and return a markdown ## Images block.
    Images smaller than 2 KB (icons, bullets) are skipped.
    Returns an empty string if nothing was transcribed.
    """
    from .llm_client import get_llm_client
    from .token_tracker import log_call
    from ..config import MODEL

    llm = get_llm_client()
    parts: list[str] = []
    total_in = total_out = 0

    for i, (label, blob) in enumerate(image_blobs[:MAX_EMBEDDED_IMAGES]):
        if len(blob) < 2048:
            continue
        try:
            log.info("  Transcribing embedded image %d: %s (%d bytes)", i + 1, label or "unlabelled", len(blob))
            description, usage = llm.transcribe_page(blob, i + 1, None)
            total_in  += usage.input_tokens
            total_out += usage.output_tokens
            parts.append(f"[Image: {label}]\n{description}" if label else f"[Image]\n{description}")
        except Exception as exc:
            log.warning("  Could not transcribe image %s: %s", label, exc)
            if label:
                parts.append(f"[Image: {label}]")

    if parts:
        log_call(operation, source_name, total_in, total_out, model=MODEL)
        log.info("Embedded image transcription done: %d images  in=%d  out=%d", len(parts), total_in, total_out)
        return "\n\n## Images\n\n" + "\n\n".join(parts)
    return ""


def extract_docx(path: Path) -> str:
    """Extract text from a Word .docx file, preserving headings, tables, and embedded images."""
    from docx import Document

    doc = Document(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    base_text = "\n\n".join(lines)

    # Extract embedded images via relationship map
    image_blobs: list[tuple[str, bytes]] = []
    seen: set[str] = set()
    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype.lower():
            continue
        if rel_id in seen:
            continue
        seen.add(rel_id)
        try:
            blob = rel.target_part.blob
            image_blobs.append((rel.target_part.partname.split("/")[-1], blob))
        except Exception as exc:
            log.warning("  Could not read DOCX image rel %s: %s", rel_id, exc)

    return base_text + _transcribe_embedded_images(image_blobs, path.name, "extract_docx_images")


def extract_markdown(path: Path) -> str:
    """Extract text from a Markdown file, vision-transcribing locally referenced images."""
    import re

    text = path.read_text(encoding="utf-8")

    # Collect local image references: ![alt](path) — skip http(s) and data URIs
    img_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    image_blobs: list[tuple[str, bytes]] = []
    for m in img_pattern.finditer(text):
        alt, src = m.group(1).strip(), m.group(2).strip()
        if src.startswith(("http://", "https://", "data:")):
            continue
        img_path = (path.parent / src).resolve()
        if not img_path.exists():
            continue
        try:
            image_blobs.append((alt or img_path.name, img_path.read_bytes()))
        except Exception as exc:
            log.warning("  Could not read markdown image %s: %s", img_path, exc)

    return text + _transcribe_embedded_images(image_blobs, path.name, "extract_markdown_images")

def extract_pubmed(pmid: str) -> dict:
    """Fetch abstract + metadata from NCBI E-utilities. Returns dict with title, authors, abstract, citation."""
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

    summary_url = f"{base}/esummary.fcgi?db=pubmed&id={pmid}&retmode=json"
    abstract_url = f"{base}/efetch.fcgi?db=pubmed&id={pmid}&rettype=abstract&retmode=text"

    with httpx.Client(timeout=30) as client:
        summary = client.get(summary_url).json()
        abstract_text = client.get(abstract_url).text

    doc_summary = summary.get("result", {}).get(pmid, {})
    title = doc_summary.get("title", "Unknown Title")
    authors = ", ".join(a.get("name", "") for a in doc_summary.get("authors", [])[:5])
    pub_date = doc_summary.get("pubdate", "")
    journal = doc_summary.get("source", "")

    citation = f"{authors}. {title}. {journal}. {pub_date}. PMID: {pmid}"

    return {
        "title": title,
        "citation": citation,
        "text": abstract_text,
        "pmid": pmid,
    }

def extract_url_browserbase(url: str) -> dict:
    """
    Fetch a JS-rendered web page via Browserbase + Playwright.

    Stagehand creates a remote Browserbase browser session; Playwright connects
    over CDP to get the fully-rendered HTML; our existing LLM pipeline then
    extracts clean text — same as the httpx path but with a real browser.
    """
    from stagehand import Stagehand
    from playwright.sync_api import sync_playwright
    from bs4 import BeautifulSoup
    from .llm_client import get_llm_client
    from .token_tracker import log_call
    from ..config import BROWSERBASE_API_KEY, BROWSERBASE_PROJECT_ID, ANTHROPIC_API_KEY, MODEL

    log.info("Browserbase fetch: %s", url)

    # 1. Start a Browserbase session via Stagehand to get the CDP WebSocket URL
    stagehand = Stagehand(
        browserbase_api_key=BROWSERBASE_API_KEY,
        browserbase_project_id=BROWSERBASE_PROJECT_ID,
        model_api_key=ANTHROPIC_API_KEY,
        server="remote",
    )
    session_resp = stagehand.sessions.start(model_name="gpt-4o")
    session_id = session_resp.data.session_id
    cdp_url = session_resp.data.cdp_url
    log.info("Browserbase session: %s", session_id)

    try:
        # 2. Connect Playwright to the remote browser over CDP and get rendered HTML
        with sync_playwright() as pw:
            browser = pw.chromium.connect_over_cdp(cdp_url)
            page = browser.contexts[0].pages[0] if browser.contexts[0].pages else browser.contexts[0].new_page()
            page.goto(url, wait_until="networkidle", timeout=30000)
            title = page.title() or urlparse(url).path
            html = page.content()
            final_url = page.url
            browser.close()

        log.info("Browserbase rendered %d bytes for %r", len(html), title)

        # 3. Strip noise tags and pass HTML to our LLM — same pipeline as extract_url
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup.find_all(["script", "style", "noscript"]):
            tag.decompose()
        clean_html = str(soup)

        EXTRACT_TOOL = {
            "name": "extracted_page_text",
            "description": "Return the full extracted text of the web page.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": (
                            "All substantive content from the page as clean plain text. "
                            "Preserve all headings (use # ## ### markdown), paragraphs, "
                            "lists, tables. Remove navigation menus, cookie banners, "
                            "ads, and other chrome."
                        )
                    }
                },
                "required": ["text"]
            }
        }

        llm = get_llm_client()
        extract_response = llm.create_message(
            messages=[{
                "role": "user",
                "content": (
                    f"Extract the COMPLETE text content from this HTML page. "
                    f"URL: {final_url}\n\n"
                    f"Include every section, every paragraph, all data tables. "
                    f"Remove only navigation menus, ads, cookie notices, and footer boilerplate.\n\n"
                    f"HTML (may be truncated):\n{clean_html[:150000]}"
                )
            }],
            tools=[EXTRACT_TOOL],
            system="You are a precise HTML-to-text extractor. Never summarise. Extract everything.",
            max_tokens=8000,
        )

        extract_block = next(
            (b for b in extract_response.content if b.type == "tool_use"), None
        )
        full_text = extract_block.input.get("text", "") if extract_block else soup.get_text(separator="\n")
        log_call("extract_url_browserbase", title,
                 extract_response.usage.input_tokens, extract_response.usage.output_tokens, model=MODEL)
        log.info("Browserbase LLM extraction done: %d chars", len(full_text))

    finally:
        try:
            stagehand.sessions.end(session_id)
        except Exception:
            pass
        stagehand.close()

    citation = f"{title}. Retrieved from: {final_url}"
    return {"title": title, "text": full_text, "citation": citation, "url": final_url}


def extract_url(url: str) -> dict:
    """
    Fetch a web page and return its full text content plus image descriptions.

    Strategy:
    1. Fetch the raw HTML with httpx.
    2. Strip only <script>/<style> noise with BeautifulSoup, then pass the
       remaining HTML to the LLM to extract clean prose.  This handles any
       page structure — server-rendered or JS-hydrated skeletons alike, because
       the LLM reads the underlying HTML rather than the rendered DOM.
    3. Download content images (>2 KB) and vision-transcribe each one;
       descriptions are appended after the main text.

    Returns dict: { title, text, citation, url }
    """
    from bs4 import BeautifulSoup
    from .llm_client import get_llm_client
    from .token_tracker import log_call
    from ..config import MODEL

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }

    log.info("Fetching URL: %s", url)
    try:
        with httpx.Client(timeout=30, follow_redirects=True, headers=headers) as client:
            resp = client.get(url)
            resp.raise_for_status()
            final_url = str(resp.url)
    except Exception as exc:
        log.warning("httpx fetch failed (%s) — falling back to Browserbase", exc)
        return extract_url_browserbase(url)

    # If the server returned a PDF, delegate to vision extraction
    content_type = resp.headers.get("content-type", "")
    if "application/pdf" in content_type or final_url.lower().endswith(".pdf"):
        import tempfile, os
        log.info("URL returned PDF — routing to vision extractor")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(resp.content)
            tmp_path = Path(tmp.name)
        try:
            text = extract_pdf_vision(tmp_path)
        finally:
            os.unlink(tmp_path)
        parsed_title = urlparse(final_url).path.split("/")[-1] or final_url
        return {
            "title": parsed_title,
            "text": text,
            "citation": f"{parsed_title}. Retrieved from: {final_url}",
            "url": final_url,
        }

    html = resp.text

    soup = BeautifulSoup(html, "html.parser")

    # Page title
    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else urlparse(final_url).path

    # Collect image URLs before stripping tags
    img_urls = []
    for img in soup.find_all("img"):
        src = img.get("src", "")
        alt = img.get("alt", "")
        if src and not src.startswith("data:"):
            abs_src = urljoin(final_url, src) if not src.startswith("http") else src
            img_urls.append((abs_src, alt))

    # Strip only noise tags — preserve all structural/content tags
    for tag in soup.find_all(["script", "style", "noscript"]):
        tag.decompose()

    clean_html = str(soup)

    # ── Step 1: LLM extracts full text from the HTML ──────────────────────────
    llm = get_llm_client()
    total_in = total_out = 0

    EXTRACT_TOOL = {
        "name": "extracted_page_text",
        "description": "Return the full extracted text of the web page.",
        "input_schema": {
            "type": "object",
            "properties": {
                "text": {
                    "type": "string",
                    "description": (
                        "All substantive content from the page as clean plain text. "
                        "Preserve all headings (use # ## ### markdown), paragraphs, "
                        "lists, tables. Remove navigation menus, cookie banners, "
                        "ads, and other chrome. Include EVERYTHING else — every "
                        "section, every paragraph, every data point."
                    )
                }
            },
            "required": ["text"]
        }
    }

    extract_response = llm.create_message(
        messages=[{
            "role": "user",
            "content": (
                f"Extract the COMPLETE text content from this HTML page. "
                f"URL: {final_url}\n\n"
                f"Include every section, every paragraph, all drug doses, "
                f"all data tables, all clinical details. Remove only navigation "
                f"menus, ads, cookie notices, and footer boilerplate.\n\n"
                f"HTML (may be truncated):\n{clean_html[:150000]}"
            )
        }],
        tools=[EXTRACT_TOOL],
        system="You are a precise HTML-to-text extractor. Never summarise. Extract everything.",
        max_tokens=8000,
    )

    total_in  += extract_response.usage.input_tokens
    total_out += extract_response.usage.output_tokens

    extract_block = next(
        (b for b in extract_response.content if b.type == "tool_use"), None
    )
    full_text = extract_block.input.get("text", "") if extract_block else soup.get_text(separator="\n")
    log.info(
        "LLM extraction: in=%d out=%d → %d chars of text",
        extract_response.usage.input_tokens,
        extract_response.usage.output_tokens,
        len(full_text),
    )

    log_call("extract_url_text", title, total_in, total_out, model=MODEL)

    # ── Step 2: Download and vision-transcribe content images ─────────────────
    image_parts = []
    image_count = 0
    img_in = img_out = 0
    MAX_IMAGES = 15

    for img_url, alt in img_urls:
        if image_count >= MAX_IMAGES:
            break
        try:
            with httpx.Client(timeout=15, follow_redirects=True, headers=headers) as client:
                img_resp = client.get(img_url)
                img_resp.raise_for_status()
                if not img_resp.headers.get("content-type", "").startswith("image/"):
                    continue
                image_bytes = img_resp.content

            if len(image_bytes) < 2048:   # skip icons / tracking pixels
                continue

            image_count += 1
            log.info("  Transcribing image %d: %s (%d bytes)", image_count, img_url, len(image_bytes))
            description, usage = llm.transcribe_page(image_bytes, image_count, None)
            img_in  += usage.input_tokens
            img_out += usage.output_tokens

            caption = f"[Image: {alt}]\n{description}" if alt else f"[Image]\n{description}"
            image_parts.append(caption)

        except Exception as exc:
            log.warning("  Could not fetch/transcribe image %s: %s", img_url, exc)
            if alt:
                image_parts.append(f"[Image: {alt}]")

    if image_count > 0:
        log_call("extract_url_images", title, img_in, img_out, model=MODEL)
        log.info("Image transcription done: %d images  in=%d  out=%d", image_count, img_in, img_out)
        full_text += "\n\n## Images\n\n" + "\n\n".join(image_parts)

    log.info("URL extraction done: %d chars total  %d images", len(full_text), image_count)

    citation = f"{title}. Retrieved from: {final_url}"
    return {
        "title": title,
        "text": full_text,
        "citation": citation,
        "url": final_url,
    }


def extract_text(path: Path, cache_dir: "Path | None" = None) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_vision(path, cache_dir=cache_dir)
    if suffix == ".docx":
        return extract_docx(path)
    return path.read_text(encoding="utf-8", errors="replace")
